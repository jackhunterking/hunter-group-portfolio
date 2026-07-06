#!/usr/bin/env python3
"""Build a Tara Housing clinical vignette DOCX."""

from __future__ import annotations

import argparse
import datetime as dt
import re
import sys
from pathlib import Path


PROHIBITED_PATTERNS = [
    r"\breported\b",
    r"\breport\b",
    r"\bfiled a report\b",
    r"\bcalled the police\b",
    r"\bcontacted authorities\b",
    r"\brestraining order\b",
    r"\bcriminal charges\b",
    r"\bcourt order\b",
    r"\b911\b",
]


def word_count(text: str) -> int:
    return len(re.findall(r"\b[\w']+\b", text))


def find_prohibited_terms(text: str) -> list[str]:
    found: list[str] = []
    for pattern in PROHIBITED_PATTERNS:
        if re.search(pattern, text, flags=re.IGNORECASE):
            found.append(pattern)
    return found


def safe_filename(value: str) -> str:
    return re.sub(r"[^A-Za-z0-9 ._-]+", "", value).strip()


def build_docx(subject: str, body: str, output_dir: Path, skip_checks: bool) -> Path:
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Pt
    except ImportError as exc:
        raise SystemExit(
            "python-docx is required. Install it with: python3 -m pip install python-docx"
        ) from exc

    count = word_count(body)
    prohibited = find_prohibited_terms(body)

    if not skip_checks:
        if count < 1500 or count > 1800:
            print(f"Warning: body word count is {count}, expected 1500-1800.", file=sys.stderr)
        if prohibited:
            print("Warning: prohibited police-trace language detected:", file=sys.stderr)
            for term in prohibited:
                print(f"  - {term}", file=sys.stderr)

    today = dt.date.today().isoformat()
    filename = f"Vignette - {safe_filename(subject)} - {today}.docx"
    output_dir.mkdir(parents=True, exist_ok=True)
    output_path = output_dir / filename

    document = Document()
    styles = document.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(11)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run(f"Clinical Observation - {subject}")
    title_run.bold = True
    title_run.font.name = "Calibri"
    title_run.font.size = Pt(11)

    date_line = document.add_paragraph()
    date_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_run = date_line.add_run(today)
    date_run.italic = True
    date_run.font.name = "Calibri"
    date_run.font.size = Pt(11)

    for raw_para in re.split(r"\n\s*\n", body.strip()):
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.first_line_indent = Pt(18)
        paragraph.add_run(raw_para.strip())

    document.add_paragraph()
    sig = document.add_paragraph()
    sig.add_run("Maryam Ashkan, RP\nRP #004730\nCRPO")

    document.save(output_path)
    return output_path


def main() -> int:
    parser = argparse.ArgumentParser(description="Build a clinical vignette DOCX.")
    parser.add_argument("--subject", required=True, help="Subject full name.")
    parser.add_argument("--body-file", required=True, type=Path, help="Plain text body file.")
    parser.add_argument("--output-dir", required=True, type=Path, help="Output directory.")
    parser.add_argument("--skip-checks", action="store_true", help="Suppress validation warnings.")
    args = parser.parse_args()

    body = args.body_file.read_text(encoding="utf-8")
    output_path = build_docx(args.subject, body, args.output_dir, args.skip_checks)
    print(output_path)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
